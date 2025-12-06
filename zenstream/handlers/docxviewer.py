import streamlit as st
from styler import Styler
from docx import Document
import re
import pandas as pd


class DocXViewer:
    @staticmethod
    def show_file(file_name: str):
        st.header("DocX Viewer")
        document = Document(file_name)

        # Show paragraphs
        paragraphs = []
        for index, paragraph in enumerate(document.paragraphs):
            paragraphs.append(
                (f"Paragraph {index}", DocXViewer._format_text(paragraph.text))
            )
        df = pd.DataFrame(paragraphs)
        with st.expander("Paragraphs"):
            st.data_editor(df, num_rows="dynamic")

        with st.expander("Tables"):
            # Show tables
            for index, table in enumerate(document.tables):
                DocXViewer._show_table(table, index + 1)

        # Show metadata
        with st.expander("Metadata"):
            Styler.show_dict(
                "Metadata",
                {
                    "Author": document.core_properties.author,
                    "Category": document.core_properties.category,
                    "Comments": document.core_properties.comments,
                    "Status": document.core_properties.content_status,
                    "Created": document.core_properties.created,
                    "Identifier": document.core_properties.identifier,
                    "Keywords": document.core_properties.keywords,
                    "Language": document.core_properties.language,
                    "Last modified by": document.core_properties.last_modified_by,
                    "Revision": document.core_properties.revision,
                    "Subject": document.core_properties.subject,
                    "Title": document.core_properties.title,
                    "Version": document.core_properties.version,
                },
            )

    @staticmethod
    def _show_table(table, index: int):
        table_rows = DocXViewer._get_table_rows(table)
        df = pd.DataFrame(table_rows)
        st.subheader(f"Table {index}")
        st.data_editor(df, num_rows="dynamic")

        # Try adding some graphs
        with st.expander(f"Table {index} graphs"):
            st.markdown(f"Table {index} analysis")
            col1, col2 = st.columns(2)
            st.divider()
            with col1:
                st.markdown("**Area Chart**")
                st.area_chart(df)
                st.markdown("**Bar Chart**")
                st.bar_chart(df)
            with col2:
                st.markdown("**Line Chart**")
                st.line_chart(df)

    @staticmethod
    def _format_text(text: str) -> str:
        """Remove non-readable characters"""
        control_char_pattern = r"[\x00-\x1F\x7F]"
        return re.sub(control_char_pattern, "", text)

    @staticmethod
    def _get_table_rows(table) -> list:
        """Return a list of lists / tuples"""
        # [
        # ("",  "a", "b"),
        # ("c", "d", ""),
        # ("",  "e", ""),
        # ]

        def iter_row_cell_texts(row):  # -> Iterator[str]:
            for _ in range(row.grid_cols_before):
                yield ""
            for c in row.cells:
                yield c.text
            for _ in range(row.grid_cols_after):
                yield ""

        rows = [tuple(iter_row_cell_texts(r)) for r in table.rows]
        return rows
